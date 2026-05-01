"""Unit tests for src/export/renderers/docx.py.

Covers the DOCX renderer's internal helpers and the public render() function:
  - _parse_paragraph: heading, bullet, numbered, plain text detection
  - _add_formatted_run: bold, italic, inline code, link inline formatting
  - _process_markdown: code blocks, mixed paragraph types
  - render(): metadata properties, title, Q&A section, empty Q&A, output bytes
"""

from __future__ import annotations

from io import BytesIO

from docx import Document

from src.export.models import DocumentParts, MetadataHeader, QAPair
from src.export.renderers.docx import (
    _add_formatted_run,
    _parse_paragraph,
    _process_markdown,
    render,
)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _make_parts(
    title: str = "Test Report",
    session_id: str = "aaaaaaaa-bbbb-cccc-dddd-eeeeeeeeeeee",
    report_body: str = "# Introduction\n\nHello world.",
    qa_pairs: list[QAPair] | None = None,
) -> DocumentParts:
    return DocumentParts(
        metadata=MetadataHeader(
            title=title,
            session_id=session_id,
            exported_at="2025-01-01T00:00:00+00:00",
            format="docx",
        ),
        report_body=report_body,
        qa_pairs=qa_pairs or [],
    )


def _load_doc(raw: bytes) -> Document:
    """Load a Document from raw bytes returned by render()."""
    return Document(BytesIO(raw))


def _all_text(doc: Document) -> str:
    """Concatenate all paragraph text in the document."""
    return "\n".join(p.text for p in doc.paragraphs)


# ---------------------------------------------------------------------------
# _parse_paragraph
# ---------------------------------------------------------------------------


class TestParseParagraph:
    def test_heading_level_1(self):
        ptype, content, level = _parse_paragraph("# My Heading")
        assert ptype == "heading"
        assert content == "My Heading"
        assert level == "1"

    def test_heading_level_2(self):
        ptype, content, level = _parse_paragraph("## Sub Heading")
        assert ptype == "heading"
        assert content == "Sub Heading"
        assert level == "2"

    def test_heading_level_6(self):
        ptype, content, level = _parse_paragraph("###### Deep Heading")
        assert ptype == "heading"
        assert content == "Deep Heading"
        assert level == "6"

    def test_bullet_dash(self):
        ptype, content, level = _parse_paragraph("- Item one")
        assert ptype == "bullet"
        assert content == "Item one"
        assert level == ""

    def test_bullet_asterisk(self):
        ptype, content, level = _parse_paragraph("* Item two")
        assert ptype == "bullet"
        assert content == "Item two"
        assert level == ""

    def test_numbered_list(self):
        ptype, content, level = _parse_paragraph("1. First item")
        assert ptype == "numbered"
        assert content == "First item"
        assert level == ""

    def test_numbered_list_higher_number(self):
        ptype, content, level = _parse_paragraph("42. Forty-second item")
        assert ptype == "numbered"
        assert content == "Forty-second item"
        assert level == ""

    def test_plain_text(self):
        ptype, content, level = _parse_paragraph("Just a plain sentence.")
        assert ptype == "text"
        assert content == "Just a plain sentence."
        assert level == ""

    def test_empty_string(self):
        ptype, content, level = _parse_paragraph("")
        assert ptype == "text"
        assert content == ""
        assert level == ""

    def test_hash_without_space_is_plain_text(self):
        # "#NoSpace" is not a valid markdown heading (no space after #)
        ptype, content, level = _parse_paragraph("#NoSpace")
        assert ptype == "text"

    def test_heading_preserves_content_with_special_chars(self):
        ptype, content, level = _parse_paragraph("## Results & Analysis (2025)")
        assert ptype == "heading"
        assert content == "Results & Analysis (2025)"


# ---------------------------------------------------------------------------
# _add_formatted_run
# ---------------------------------------------------------------------------


class TestAddFormattedRun:
    def _para(self):
        """Return a fresh paragraph from a temporary Document."""
        return Document().add_paragraph()

    def test_plain_text_added_as_single_run(self):
        para = self._para()
        _add_formatted_run(para, "Hello world")
        text = "".join(r.text for r in para.runs)
        assert text == "Hello world"

    def test_bold_double_asterisk(self):
        para = self._para()
        _add_formatted_run(para, "**bold text**")
        bold_runs = [r for r in para.runs if r.bold]
        assert any(r.text == "bold text" for r in bold_runs)

    def test_bold_double_underscore(self):
        para = self._para()
        _add_formatted_run(para, "__bold text__")
        bold_runs = [r for r in para.runs if r.bold]
        assert any(r.text == "bold text" for r in bold_runs)

    def test_italic_single_asterisk(self):
        para = self._para()
        _add_formatted_run(para, "*italic text*")
        italic_runs = [r for r in para.runs if r.italic]
        assert any(r.text == "italic text" for r in italic_runs)

    def test_italic_single_underscore(self):
        para = self._para()
        _add_formatted_run(para, "_italic text_")
        italic_runs = [r for r in para.runs if r.italic]
        assert any(r.text == "italic text" for r in italic_runs)

    def test_inline_code_uses_courier_font(self):
        para = self._para()
        _add_formatted_run(para, "`some_code()`")
        code_runs = [r for r in para.runs if r.font.name == "Courier New"]
        assert any(r.text == "some_code()" for r in code_runs)

    def test_link_renders_text_and_url(self):
        para = self._para()
        _add_formatted_run(para, "[Click here](https://example.com)")
        full_text = "".join(r.text for r in para.runs)
        assert "Click here" in full_text
        assert "https://example.com" in full_text

    def test_link_run_is_underlined(self):
        para = self._para()
        _add_formatted_run(para, "[Link](https://example.com)")
        underlined = [r for r in para.runs if r.underline]
        assert len(underlined) >= 1

    def test_mixed_bold_and_plain(self):
        para = self._para()
        _add_formatted_run(para, "Start **bold** end")
        full_text = "".join(r.text for r in para.runs)
        assert "Start " in full_text
        assert "bold" in full_text
        assert "end" in full_text

    def test_empty_string_produces_no_runs_or_empty_run(self):
        para = self._para()
        _add_formatted_run(para, "")
        # Either no runs or a single empty run — no crash
        full_text = "".join(r.text for r in para.runs)
        assert full_text == ""


# ---------------------------------------------------------------------------
# _process_markdown
# ---------------------------------------------------------------------------


class TestProcessMarkdown:
    def _doc_with_md(self, md: str) -> Document:
        doc = Document()
        _process_markdown(doc, md)
        return doc

    def test_heading_creates_heading_paragraph(self):
        doc = self._doc_with_md("# Section One")
        texts = [p.text for p in doc.paragraphs]
        assert "Section One" in texts

    def test_bullet_list_items_present(self):
        md = "- Alpha\n- Beta\n- Gamma"
        doc = self._doc_with_md(md)
        texts = _all_text(doc)
        assert "Alpha" in texts
        assert "Beta" in texts
        assert "Gamma" in texts

    def test_numbered_list_items_present(self):
        md = "1. First\n2. Second\n3. Third"
        doc = self._doc_with_md(md)
        texts = _all_text(doc)
        assert "First" in texts
        assert "Second" in texts
        assert "Third" in texts

    def test_code_block_uses_courier_font(self):
        md = "```python\nprint('hello')\n```"
        doc = self._doc_with_md(md)
        courier_runs = [
            r for p in doc.paragraphs for r in p.runs if r.font.name == "Courier New"
        ]
        assert len(courier_runs) > 0

    def test_code_block_language_label_present(self):
        md = "```python\nprint('hello')\n```"
        doc = self._doc_with_md(md)
        all_run_text = " ".join(r.text for p in doc.paragraphs for r in p.runs)
        assert "python" in all_run_text

    def test_code_block_without_language(self):
        md = "```\nsome code\n```"
        doc = self._doc_with_md(md)
        courier_runs = [
            r for p in doc.paragraphs for r in p.runs if r.font.name == "Courier New"
        ]
        assert any("some code" in r.text for r in courier_runs)

    def test_empty_markdown_produces_no_paragraphs(self):
        doc = self._doc_with_md("")
        # Document() always has a default empty paragraph; extra ones should not appear
        non_empty = [p for p in doc.paragraphs if p.text.strip()]
        assert len(non_empty) == 0

    def test_multiple_paragraph_blocks(self):
        md = "First paragraph.\n\nSecond paragraph."
        doc = self._doc_with_md(md)
        texts = _all_text(doc)
        assert "First paragraph." in texts
        assert "Second paragraph." in texts

    def test_blank_lines_between_items_ignored(self):
        md = "- Item A\n\n- Item B"
        doc = self._doc_with_md(md)
        texts = _all_text(doc)
        assert "Item A" in texts
        assert "Item B" in texts


# ---------------------------------------------------------------------------
# render() — output bytes and structure
# ---------------------------------------------------------------------------


class TestRender:
    def test_returns_bytes(self):
        result = render(_make_parts())
        assert isinstance(result, bytes)

    def test_returns_non_empty_bytes(self):
        result = render(_make_parts())
        assert len(result) > 0

    def test_output_is_valid_docx(self):
        """Bytes must be parseable as a DOCX (ZIP-based) document."""
        result = render(_make_parts())
        doc = _load_doc(result)
        assert doc is not None

    def test_output_starts_with_zip_magic_bytes(self):
        """DOCX files are ZIP archives; first two bytes are PK (0x50 0x4B)."""
        result = render(_make_parts())
        assert result[:2] == b"PK"

    def test_core_properties_title(self):
        result = render(_make_parts(title="My Research Report"))
        doc = _load_doc(result)
        assert doc.core_properties.title == "My Research Report"

    def test_core_properties_author(self):
        result = render(_make_parts())
        doc = _load_doc(result)
        assert doc.core_properties.author == "Deep Research"

    def test_core_properties_subject_contains_session_id(self):
        sid = "aaaaaaaa-bbbb-cccc-dddd-eeeeeeeeeeee"
        result = render(_make_parts(session_id=sid))
        doc = _load_doc(result)
        assert sid in doc.core_properties.subject

    def test_title_paragraph_present(self):
        result = render(_make_parts(title="Unique Title XYZ"))
        doc = _load_doc(result)
        texts = _all_text(doc)
        assert "Unique Title XYZ" in texts

    def test_session_id_in_document(self):
        sid = "aaaaaaaa-bbbb-cccc-dddd-eeeeeeeeeeee"
        result = render(_make_parts(session_id=sid))
        doc = _load_doc(result)
        texts = _all_text(doc)
        assert sid in texts

    def test_exported_at_in_document(self):
        result = render(_make_parts())
        doc = _load_doc(result)
        texts = _all_text(doc)
        assert "2025-01-01T00:00:00+00:00" in texts

    def test_report_body_content_present(self):
        result = render(_make_parts(report_body="## Findings\n\nSome findings here."))
        doc = _load_doc(result)
        texts = _all_text(doc)
        assert "Findings" in texts
        assert "Some findings here." in texts

    def test_qa_section_heading_present_when_pairs_exist(self):
        pairs = [QAPair(question="What is X?", answer="X is Y.")]
        result = render(_make_parts(qa_pairs=pairs))
        doc = _load_doc(result)
        texts = _all_text(doc)
        assert "Q&A History" in texts

    def test_qa_question_present(self):
        pairs = [QAPair(question="What is X?", answer="X is Y.")]
        result = render(_make_parts(qa_pairs=pairs))
        doc = _load_doc(result)
        texts = _all_text(doc)
        assert "What is X?" in texts

    def test_qa_answer_present(self):
        pairs = [QAPair(question="What is X?", answer="X is Y.")]
        result = render(_make_parts(qa_pairs=pairs))
        doc = _load_doc(result)
        texts = _all_text(doc)
        assert "X is Y." in texts

    def test_multiple_qa_pairs_all_present(self):
        pairs = [
            QAPair(question="Q1", answer="A1"),
            QAPair(question="Q2", answer="A2"),
            QAPair(question="Q3", answer="A3"),
        ]
        result = render(_make_parts(qa_pairs=pairs))
        doc = _load_doc(result)
        texts = _all_text(doc)
        for i in range(1, 4):
            assert f"Q{i}" in texts
            assert f"A{i}" in texts

    def test_no_qa_section_when_pairs_empty(self):
        result = render(_make_parts(qa_pairs=[]))
        doc = _load_doc(result)
        texts = _all_text(doc)
        assert "Q&A History" not in texts

    def test_empty_report_body_does_not_crash(self):
        result = render(_make_parts(report_body=""))
        assert isinstance(result, bytes)
        assert len(result) > 0

    def test_render_is_deterministic(self):
        """Two calls with identical inputs must produce identical bytes."""
        parts = _make_parts(
            report_body="## Section\n\nContent.",
            qa_pairs=[QAPair(question="Q?", answer="A.")],
        )
        first = render(parts)
        second = render(parts)
        assert first == second

    def test_report_body_with_heading_levels(self):
        body = "# H1\n\n## H2\n\n### H3\n\nParagraph text."
        result = render(_make_parts(report_body=body))
        doc = _load_doc(result)
        texts = _all_text(doc)
        assert "H1" in texts
        assert "H2" in texts
        assert "H3" in texts
        assert "Paragraph text." in texts

    def test_report_body_with_bullet_list(self):
        body = "- Alpha\n- Beta\n- Gamma"
        result = render(_make_parts(report_body=body))
        doc = _load_doc(result)
        texts = _all_text(doc)
        assert "Alpha" in texts
        assert "Beta" in texts
        assert "Gamma" in texts

    def test_report_body_with_numbered_list(self):
        body = "1. First\n2. Second\n3. Third"
        result = render(_make_parts(report_body=body))
        doc = _load_doc(result)
        texts = _all_text(doc)
        assert "First" in texts
        assert "Second" in texts
        assert "Third" in texts

    def test_report_body_with_code_block(self):
        body = "```python\nprint('hello')\n```"
        result = render(_make_parts(report_body=body))
        doc = _load_doc(result)
        courier_runs = [
            r for p in doc.paragraphs for r in p.runs if r.font.name == "Courier New"
        ]
        assert len(courier_runs) > 0

    def test_report_body_with_bold_inline(self):
        body = "This is **important** text."
        result = render(_make_parts(report_body=body))
        doc = _load_doc(result)
        bold_runs = [r for p in doc.paragraphs for r in p.runs if r.bold]
        assert any("important" in r.text for r in bold_runs)

    def test_qa_pair_with_empty_answer(self):
        """A QAPair with an empty answer must not crash the renderer."""
        pairs = [QAPair(question="Unanswered?", answer="")]
        result = render(_make_parts(qa_pairs=pairs))
        doc = _load_doc(result)
        texts = _all_text(doc)
        assert "Unanswered?" in texts

    def test_qa_pair_with_empty_question(self):
        """A QAPair with an empty question (lone assistant message) must not crash."""
        pairs = [QAPair(question="", answer="Unprompted answer.")]
        result = render(_make_parts(qa_pairs=pairs))
        doc = _load_doc(result)
        texts = _all_text(doc)
        assert "Unprompted answer." in texts

    def test_title_at_255_char_limit_does_not_crash(self):
        """python-docx core_properties.title has a 255-char limit.
        A title exactly at the limit must render without error."""
        title_255 = "A" * 255
        result = render(_make_parts(title=title_255))
        assert isinstance(result, bytes)
        assert len(result) > 0

    def test_title_over_255_chars_is_truncated_in_core_props(self):
        """Titles longer than 255 chars are truncated to 255 in core properties."""
        long_title = "B" * 400
        result = render(_make_parts(title=long_title))
        doc = _load_doc(result)
        assert doc.core_properties.title == "B" * 255

    def test_title_over_255_chars_full_title_in_body(self):
        """The full (unsanitized) title still appears in the document heading paragraph."""
        long_title = "C" * 400
        result = render(_make_parts(title=long_title))
        doc = _load_doc(result)
        texts = _all_text(doc)
        assert long_title in texts

    def test_unicode_content_in_report_body(self):
        body = "Résumé: café, naïve, über, 日本語, العربية"
        result = render(_make_parts(report_body=body))
        doc = _load_doc(result)
        texts = _all_text(doc)
        assert "café" in texts
        assert "日本語" in texts


# ---------------------------------------------------------------------------
# _sanitize helper
# ---------------------------------------------------------------------------


class TestSanitize:
    """Tests for the _sanitize() and _core_prop() helpers."""

    def test_sanitize_removes_null_byte(self):
        from src.export.renderers.docx import _sanitize

        assert _sanitize("hello\x00world") == "helloworld"

    def test_sanitize_removes_c0_control_chars(self):
        from src.export.renderers.docx import _sanitize

        # \x01–\x08, \x0b, \x0c, \x0e–\x1f are illegal
        assert _sanitize("a\x01b\x08c\x0bd\x1fe") == "abcde"

    def test_sanitize_preserves_tab_newline_cr(self):
        from src.export.renderers.docx import _sanitize

        text = "line1\nline2\ttabbed\rcarriage"
        assert _sanitize(text) == text

    def test_sanitize_removes_del_and_c1_chars(self):
        from src.export.renderers.docx import _sanitize

        assert _sanitize("a\x7fb\x80c\x9fd") == "abcd"

    def test_sanitize_passes_through_normal_unicode(self):
        from src.export.renderers.docx import _sanitize

        text = "café 日本語 العربية"
        assert _sanitize(text) == text

    def test_sanitize_empty_string(self):
        from src.export.renderers.docx import _sanitize

        assert _sanitize("") == ""

    def test_core_prop_truncates_to_255(self):
        from src.export.renderers.docx import _core_prop

        assert _core_prop("x" * 300) == "x" * 255

    def test_core_prop_strips_control_chars_then_truncates(self):
        from src.export.renderers.docx import _core_prop

        # 200 normal chars + 100 null bytes → sanitized to 200 chars (no truncation needed)
        assert _core_prop("a" * 200 + "\x00" * 100) == "a" * 200

    def test_render_with_control_chars_in_report_body_does_not_crash(self):
        """Control characters in report_body must be stripped, not crash."""
        body = "Normal text\x00with\x01null\x1fbytes."
        result = render(_make_parts(report_body=body))
        assert isinstance(result, bytes)
        doc = _load_doc(result)
        texts = _all_text(doc)
        assert "Normal text" in texts
        assert "withbytes." in texts or "with" in texts

    def test_render_with_control_chars_in_qa_does_not_crash(self):
        """Control characters in Q&A content must be stripped, not crash."""
        pairs = [QAPair(question="Q\x00uestion?", answer="Ans\x1fwer.")]
        result = render(_make_parts(qa_pairs=pairs))
        assert isinstance(result, bytes)
        doc = _load_doc(result)
        texts = _all_text(doc)
        assert "uestion?" in texts
        assert "Ans" in texts

    def test_render_with_control_chars_in_title_does_not_crash(self):
        """Control characters in title must be stripped, not crash."""
        result = render(_make_parts(title="Title\x00With\x01Control"))
        assert isinstance(result, bytes)
        doc = _load_doc(result)
        assert doc.core_properties.title == "TitleWithControl"
