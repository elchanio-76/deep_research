"""Property-based tests for the export renderers.

Feature: export-formats
Uses hypothesis @given with @settings(max_examples=100).
"""

from hypothesis import given, settings
from hypothesis import strategies as st

from src.export.models import DocumentParts, MetadataHeader, QAPair
from src.export.renderers import markdown as markdown_renderer
from src.export.renderers import pdf as pdf_renderer
from src.export.renderers import docx as docx_renderer

import uuid as _uuid_module

import asyncpg as _asyncpg
from fastapi import FastAPI as _FastAPI
from fastapi.testclient import TestClient as _TestClient
from unittest.mock import (
    AsyncMock as _AsyncMock,
    MagicMock as _MagicMock,
    patch as _patch,
)

from src.export.errors import (
    RenderError as _RenderError,
    ReportNotReadyError as _ReportNotReadyError,
    SessionNotFoundError as _SessionNotFoundError,
)
from src.export.router import router as _export_router

from src.export.models import ExportFormat, ExportResult

# ---------------------------------------------------------------------------
# Hypothesis strategies
# ---------------------------------------------------------------------------

metadata_strategy = st.builds(
    MetadataHeader,
    title=st.text(min_size=1, max_size=200),
    session_id=st.uuids().map(str),
    exported_at=st.just("2025-01-01T00:00:00+00:00"),
    format=st.sampled_from(["markdown", "pdf", "docx"]),
)

qa_pair_strategy = st.builds(
    QAPair,
    question=st.text(min_size=1),
    answer=st.text(min_size=1),
)

document_parts_strategy = st.builds(
    DocumentParts,
    metadata=metadata_strategy,
    report_body=st.text(min_size=0, max_size=2000),
    qa_pairs=st.lists(qa_pair_strategy, max_size=20),
)


# ---------------------------------------------------------------------------
# Property 1: Metadata header completeness and document ordering
# Feature: export-formats, Property 1: metadata completeness and document ordering
# Validates: Requirements 1.3, 1.4, 3.1, 3.2
# ---------------------------------------------------------------------------


@given(parts=document_parts_strategy)
@settings(max_examples=100)
def test_property_1_metadata_completeness_and_ordering(parts: DocumentParts) -> None:
    """All four metadata fields are present and the document sections appear
    in the correct order: metadata block → report body → Q&A section."""
    output = markdown_renderer.render(parts)

    # --- Metadata field presence ---
    assert (
        f"title: {parts.metadata.title}" in output
    ), "title field missing from metadata"
    assert (
        f"session_id: {parts.metadata.session_id}" in output
    ), "session_id field missing"
    assert (
        f"exported_at: {parts.metadata.exported_at}" in output
    ), "exported_at field missing"
    assert f"format: {parts.metadata.format}" in output, "format field missing"

    # --- Section ordering ---
    # The metadata block is delimited by opening "---" and closing "---".
    first_fence = output.index("---")
    closing_fence = output.index("---", first_fence + 3)
    metadata_end = closing_fence + 3  # position just after the closing "---"

    # Report body must appear after the metadata block.
    # Search from metadata_end to avoid false positives when report_body text
    # also appears inside the metadata block (e.g. title == report_body).
    if parts.report_body:
        report_pos = output.find(parts.report_body, metadata_end)
        assert report_pos != -1, "Report body not found after the metadata block"

    # Q&A section must appear after the report body (when qa_pairs are present).
    if parts.qa_pairs:
        qa_heading_pos = output.index("## Q&A History")
        assert (
            qa_heading_pos > metadata_end
        ), "Q&A History section appears before or inside the metadata block"
        if parts.report_body:
            report_pos = output.find(parts.report_body, metadata_end)
            assert (
                qa_heading_pos > report_pos
            ), "Q&A History section appears before the report body"


# ---------------------------------------------------------------------------
# Property 2: Q&A block count invariant
# Feature: export-formats, Property 2: Q&A block count invariant
# Validates: Requirements 1.5, 7.3
# ---------------------------------------------------------------------------

# Strategy: generate a list of (role, content) message tuples where role is
# either "user" or "assistant".  The service pairs them into QAPairs by
# iterating sequentially and matching user→question / assistant→answer.
# We replicate that pairing logic here so the test is self-contained.

_message_strategy = st.tuples(
    st.sampled_from(["user", "assistant"]),
    st.text(min_size=1, max_size=200),
)


def _messages_to_qa_pairs(messages: list[tuple[str, str]]) -> list[QAPair]:
    """Pair messages into QAPairs the same way Export_Service does.

    Iterates through messages in order; each consecutive (user, assistant)
    pair forms one QAPair.  Unpaired trailing messages are included with an
    empty counterpart, matching the design doc specification.
    """
    pairs: list[QAPair] = []
    i = 0
    while i < len(messages):
        role, content = messages[i]
        if role == "user":
            if i + 1 < len(messages) and messages[i + 1][0] == "assistant":
                pairs.append(QAPair(question=content, answer=messages[i + 1][1]))
                i += 2
            else:
                pairs.append(QAPair(question=content, answer=""))
                i += 1
        else:
            # Lone assistant message with no preceding user message
            pairs.append(QAPair(question="", answer=content))
            i += 1
    return pairs


@given(
    metadata=metadata_strategy,
    report_body=st.text(min_size=0, max_size=500),
    messages=st.lists(_message_strategy, max_size=30),
)
@settings(max_examples=100)
def test_property_2_qa_block_count_invariant(
    metadata: MetadataHeader,
    report_body: str,
    messages: list[tuple[str, str]],
) -> None:
    """The number of ``**A:**`` blocks in the rendered output equals the
    number of QAPairs derived from the input message list.

    Validates: Requirements 1.5, 7.3
    """
    qa_pairs = _messages_to_qa_pairs(messages)
    parts = DocumentParts(metadata=metadata, report_body=report_body, qa_pairs=qa_pairs)
    output = markdown_renderer.render(parts)

    answer_block_count = output.count("**A:**")

    assert answer_block_count == len(qa_pairs), (
        f"Expected {len(qa_pairs)} '**A:**' blocks but found {answer_block_count}. "
        f"messages={messages!r}"
    )


# ---------------------------------------------------------------------------
# Property 3: Report body verbatim preservation
# Feature: export-formats, Property 3: report body verbatim preservation
# Validates: Requirements 7.2, 7.4
# ---------------------------------------------------------------------------

# Strategy: generate arbitrary report_markdown strings including Markdown
# syntax characters (headings, code fences, bold, links, etc.) to ensure
# the renderer never escapes or modifies the body content.
_report_body_strategy = st.text(
    alphabet=st.characters(blacklist_categories=("Cs",)),  # exclude surrogates
    min_size=0,
    max_size=2000,
)


@given(
    metadata=metadata_strategy,
    report_body=_report_body_strategy,
    qa_pairs=st.lists(qa_pair_strategy, max_size=20),
)
@settings(max_examples=100)
def test_property_3_report_body_verbatim_preservation(
    metadata: MetadataHeader,
    report_body: str,
    qa_pairs: list[QAPair],
) -> None:
    """The original report_markdown string appears verbatim (as a substring)
    in the rendered output without any escaping or modification.

    Validates: Requirements 7.2, 7.4
    """
    parts = DocumentParts(metadata=metadata, report_body=report_body, qa_pairs=qa_pairs)
    output = markdown_renderer.render(parts)

    assert report_body in output, (
        "report_body was not found verbatim in the rendered output. "
        f"report_body={report_body!r}"
    )


# ---------------------------------------------------------------------------
# Property 4: Renderer determinism
# Feature: export-formats, Property 4: renderer determinism
# Validates: Requirements 7.1
# ---------------------------------------------------------------------------


@given(parts=document_parts_strategy)
@settings(max_examples=100)
def test_property_4_renderer_determinism(parts: DocumentParts) -> None:
    """Calling render(parts) twice produces byte-for-byte identical output.

    Validates: Requirements 7.1
    """
    first = markdown_renderer.render(parts)
    second = markdown_renderer.render(parts)

    assert first == second, (
        "MarkdownRenderer.render is not deterministic: "
        "two calls with identical inputs produced different output."
    )


# ---------------------------------------------------------------------------
# Property 5: Title derivation from initial_prompt
# Feature: export-formats, Property 5: title derivation from initial_prompt
# Validates: Requirements 3.3
# ---------------------------------------------------------------------------


@given(initial_prompt=st.text(min_size=0, max_size=300))
@settings(max_examples=100)
def test_property_5_title_derivation_from_initial_prompt(initial_prompt: str) -> None:
    """When session header is None, the derived title equals initial_prompt
    when len <= 120, and initial_prompt[:120] + '…' when len > 120.

    Validates: Requirements 3.3
    """
    derived = MetadataHeader.derive_title(header=None, initial_prompt=initial_prompt)

    if len(initial_prompt) <= 120:
        assert derived == initial_prompt, (
            f"Expected title == initial_prompt for short prompt, "
            f"got {derived!r} (len={len(initial_prompt)})"
        )
    else:
        expected = initial_prompt[:120] + "\u2026"
        assert (
            derived == expected
        ), f"Expected truncated title {expected!r}, got {derived!r}"


# ---------------------------------------------------------------------------
# Property 6: Output filename pattern
# Feature: export-formats, Property 6: output filename pattern
# Validates: Requirements 4.4
# ---------------------------------------------------------------------------


@given(
    session_id=st.uuids(),
    fmt=st.sampled_from(list(ExportFormat)),
)
@settings(max_examples=100)
def test_property_6_output_filename_pattern(
    session_id: _uuid_module.UUID,
    fmt: ExportFormat,
) -> None:
    """For any session UUID and export format, the derived filename matches
    ``report-{session_id}.md`` for Markdown and ``report-{session_id}.pdf``
    for PDF.

    Validates: Requirements 4.4
    """
    filename = ExportResult.filename_for(str(session_id), fmt)

    expected_ext = (
        "md"
        if fmt == ExportFormat.markdown
        else "pdf" if fmt == ExportFormat.pdf else "docx"
    )
    expected = f"report-{session_id}.{expected_ext}"

    assert filename == expected, (
        f"Filename mismatch: expected {expected!r}, got {filename!r} "
        f"(session_id={session_id}, fmt={fmt})"
    )


# ---------------------------------------------------------------------------
# Property 8: PDF output is valid PDF bytes
# Feature: export-formats, Property 8: PDF output is valid PDF bytes
# Validates: Requirements 2.3
# ---------------------------------------------------------------------------


# Strategy restricted to printable ASCII to avoid a fontTools/weasyprint bug
# triggered by certain Unicode codepoints (OS/2 Unicode range bit > 122).
_safe_text = st.text(
    alphabet=st.characters(
        whitelist_categories=("Lu", "Ll", "Nd", "Zs", "Po", "Pd"),
        whitelist_characters=" \n",
        max_codepoint=0x007E,  # printable ASCII only
    ),
    min_size=0,
    max_size=500,
)

_safe_qa_pair = st.builds(
    QAPair,
    question=_safe_text.filter(lambda s: len(s) >= 1),
    answer=_safe_text.filter(lambda s: len(s) >= 1),
)

_safe_document_parts = st.builds(
    DocumentParts,
    metadata=metadata_strategy,
    report_body=_safe_text,
    qa_pairs=st.lists(_safe_qa_pair, max_size=5),
)


@given(parts=_safe_document_parts)
@settings(max_examples=20, deadline=None)
def test_property_8_pdf_output_is_valid_pdf_bytes(parts: DocumentParts) -> None:
    """PDFRenderer.render(parts) returns non-empty bytes whose first four
    bytes are b"%PDF", confirming a well-formed PDF was produced.

    Validates: Requirements 2.3
    """
    result = pdf_renderer.render(parts)

    assert isinstance(result, bytes), "render() must return bytes"
    assert len(result) > 0, "render() returned empty bytes"
    assert (
        result[:4] == b"%PDF"
    ), f"Output does not start with b'%PDF'; got {result[:8]!r}"


# ---------------------------------------------------------------------------
# Router property tests — require a TestClient
# ---------------------------------------------------------------------------


def _make_test_client() -> _TestClient:
    """Minimal FastAPI app with the export router for property tests."""
    app = _FastAPI()
    app.include_router(_export_router, prefix="/api")
    app.state.pool = _MagicMock(spec=_asyncpg.Pool)
    return _TestClient(app, raise_server_exceptions=False)


_VALID_UUID = str(_uuid_module.UUID("aaaaaaaa-bbbb-cccc-dddd-eeeeeeeeeeee"))

# ---------------------------------------------------------------------------
# Property 7: Invalid delivery_mode produces 422
# Feature: export-formats, Property 7: invalid delivery_mode produces 422
# Validates: Requirements 4.1, 4.5
# ---------------------------------------------------------------------------

_VALID_DELIVERY_MODES = {"download", "url"}


@given(
    delivery_mode=st.text(
        alphabet=st.characters(
            whitelist_categories=(
                "L",
                "N",
                "P",
                "S",
            ),  # letters, digits, punctuation, symbols
        ),
        min_size=1,
        max_size=50,
    ).filter(lambda s: s not in _VALID_DELIVERY_MODES)
)
@settings(max_examples=100)
def test_property_7_invalid_delivery_mode_produces_422(delivery_mode: str) -> None:
    """Any delivery_mode value that is not 'download' or 'url' must produce HTTP 422.

    Feature: export-formats, Property 7: invalid delivery_mode produces 422
    Validates: Requirements 4.1, 4.5
    """
    client = _make_test_client()
    resp = client.get(
        f"/api/export/{_VALID_UUID}/markdown?delivery_mode={delivery_mode}"
    )
    assert (
        resp.status_code == 422
    ), f"Expected 422 for delivery_mode={delivery_mode!r}, got {resp.status_code}"


# ---------------------------------------------------------------------------
# Property 9: No stack traces in error responses
# Feature: export-formats, Property 9: no stack traces in error responses
# Validates: Requirements 8.4
# ---------------------------------------------------------------------------

_STACK_TRACE_MARKERS = ["Traceback", 'File "', "/home/", "/usr/", "/opt/"]

_exception_strategy = st.one_of(
    st.just(_SessionNotFoundError("session not found")),
    st.just(_ReportNotReadyError("report not ready")),
    st.just(_RenderError("render failed")),
    st.just(_asyncpg.PostgresError()),
    st.just(RuntimeError("unexpected internal error")),
    st.just(ValueError("bad value")),
    st.just(OSError("disk full")),
)


@given(exc=_exception_strategy)
@settings(max_examples=100)
def test_property_9_no_stack_traces_in_error_responses(exc: Exception) -> None:
    """Error responses must never contain Python stack trace markers or absolute paths.

    Feature: export-formats, Property 9: no stack traces in error responses
    Validates: Requirements 8.4
    """
    client = _make_test_client()

    with _patch("src.export.router.export", new=_AsyncMock(side_effect=exc)):
        resp = client.get(f"/api/export/{_VALID_UUID}/markdown")

    body = resp.text
    for marker in _STACK_TRACE_MARKERS:
        assert marker not in body, (
            f"Stack trace marker {marker!r} found in error response body "
            f"(exc={exc!r}, status={resp.status_code}): {body!r}"
        )


# ---------------------------------------------------------------------------
# Property 10: DOCX output is valid DOCX bytes (ZIP magic bytes)
# Feature: export-formats, Property 10: DOCX output is valid DOCX bytes
# Validates: Requirements 3.x (DOCX format correctness)
# ---------------------------------------------------------------------------

from io import BytesIO as _BytesIO
from docx import Document as _DocxDocument

# python-docx / lxml requires XML-compatible text: no NULL bytes or C0/C1
# control characters.  The renderer now sanitizes input, so we restrict
# strategies only to exclude surrogates (which Python itself cannot encode).
_xml_safe_text = st.text(
    alphabet=st.characters(blacklist_categories=("Cs",)),
    min_size=0,
    max_size=500,
)

# python-docx core_properties.title has a hard 255-char limit; the renderer
# truncates after sanitization, so titles up to 500 chars are fine as input.
_xml_safe_title = st.text(
    alphabet=st.characters(blacklist_categories=("Cs",)),
    min_size=1,
    max_size=500,
)

_xml_safe_qa_pair = st.builds(
    QAPair,
    question=_xml_safe_text.filter(lambda s: len(s) >= 1),
    answer=_xml_safe_text.filter(lambda s: len(s) >= 1),
)

_xml_safe_metadata = st.builds(
    MetadataHeader,
    title=_xml_safe_title,
    session_id=st.uuids().map(str),
    exported_at=st.just("2025-01-01T00:00:00+00:00"),
    format=st.sampled_from(["markdown", "pdf", "docx"]),
)

_xml_safe_document_parts = st.builds(
    DocumentParts,
    metadata=_xml_safe_metadata,
    report_body=_xml_safe_text,
    qa_pairs=st.lists(_xml_safe_qa_pair, max_size=5),
)


@given(parts=_xml_safe_document_parts)
@settings(max_examples=50, deadline=None)
def test_property_10_docx_output_is_valid_docx_bytes(parts: DocumentParts) -> None:
    """docx_renderer.render(parts) returns non-empty bytes that start with the
    ZIP magic bytes (b'PK') and can be parsed as a valid DOCX document.

    Validates: DOCX format correctness requirement.
    Note: strategies are restricted to XML-safe characters because python-docx
    uses lxml internally and rejects C0/C1 control characters.
    """
    result = docx_renderer.render(parts)

    assert isinstance(result, bytes), "render() must return bytes"
    assert len(result) > 0, "render() returned empty bytes"
    assert (
        result[:2] == b"PK"
    ), f"DOCX output does not start with ZIP magic bytes b'PK'; got {result[:4]!r}"
    # Must be parseable as a real DOCX document
    doc = _DocxDocument(_BytesIO(result))
    assert doc is not None


# ---------------------------------------------------------------------------
# Property 11: DOCX renderer structural consistency
# Feature: export-formats, Property 11: DOCX renderer structural consistency
# Validates: Requirements 7.1 (consistent rendering)
# Note: python-docx embeds ZIP timestamps that vary between calls, so raw
# byte equality is not guaranteed.  We verify structural consistency instead:
# same paragraph count and same concatenated text on every call.
# ---------------------------------------------------------------------------


@given(parts=_xml_safe_document_parts)
@settings(max_examples=50, deadline=None)
def test_property_11_docx_renderer_structural_consistency(parts: DocumentParts) -> None:
    """Calling docx_renderer.render(parts) twice produces documents with
    identical paragraph count and identical concatenated paragraph text.

    Note: raw byte equality is not asserted because python-docx embeds
    ZIP modification timestamps that differ between calls.

    Validates: Requirements 7.1 (consistent rendering)
    """

    def _extract(raw: bytes) -> tuple[int, str]:
        doc = _DocxDocument(_BytesIO(raw))
        texts = [p.text for p in doc.paragraphs]
        return len(texts), "\n".join(texts)

    first_count, first_text = _extract(docx_renderer.render(parts))
    second_count, second_text = _extract(docx_renderer.render(parts))

    assert (
        first_count == second_count
    ), f"Paragraph count differs between calls: {first_count} vs {second_count}"
    assert (
        first_text == second_text
    ), "Concatenated paragraph text differs between calls."


# ---------------------------------------------------------------------------
# Property 12: DOCX core properties contain title and session_id
# Feature: export-formats, Property 12: DOCX metadata completeness
# Validates: Requirements 3.1, 3.2 (metadata in output)
# ---------------------------------------------------------------------------


@given(parts=_xml_safe_document_parts)
@settings(max_examples=50, deadline=None)
def test_property_12_docx_metadata_completeness(parts: DocumentParts) -> None:
    """The rendered DOCX document's core properties contain the (possibly
    truncated/sanitized) title and a subject that includes the session_id.

    The renderer strips XML-illegal control characters and truncates to 255
    chars before writing core properties, so we verify against the processed
    value rather than the raw input.

    Validates: Requirements 3.1, 3.2
    """
    import re as _re

    _illegal = _re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]")
    expected_title = _illegal.sub("", parts.metadata.title)[:255]

    result = docx_renderer.render(parts)
    doc = _DocxDocument(_BytesIO(result))

    assert doc.core_properties.title == expected_title, (
        f"core_properties.title mismatch: expected {expected_title!r}, "
        f"got {doc.core_properties.title!r}"
    )
    assert parts.metadata.session_id in doc.core_properties.subject, (
        f"session_id {parts.metadata.session_id!r} not found in "
        f"core_properties.subject {doc.core_properties.subject!r}"
    )


# ---------------------------------------------------------------------------
# Property 13: DOCX Q&A pair count invariant
# Feature: export-formats, Property 13: DOCX Q&A block count invariant
# Validates: Requirements 1.5, 7.3 (Q&A completeness)
# ---------------------------------------------------------------------------


@given(
    metadata=_xml_safe_metadata,
    report_body=_xml_safe_text,
    qa_pairs=st.lists(_xml_safe_qa_pair, max_size=10),
)
@settings(max_examples=50, deadline=None)
def test_property_13_docx_qa_pair_count_invariant(
    metadata: MetadataHeader,
    report_body: str,
    qa_pairs: list[QAPair],
) -> None:
    """The rendered DOCX document contains exactly as many 'Q: ' prefixed
    paragraphs as there are QAPairs in the input.

    Validates: Requirements 1.5, 7.3
    """
    parts = DocumentParts(metadata=metadata, report_body=report_body, qa_pairs=qa_pairs)
    result = docx_renderer.render(parts)
    doc = _DocxDocument(_BytesIO(result))

    q_paragraphs = [p for p in doc.paragraphs if any(r.text == "Q: " for r in p.runs)]

    assert len(q_paragraphs) == len(qa_pairs), (
        f"Expected {len(qa_pairs)} 'Q: ' paragraphs but found {len(q_paragraphs)}. "
        f"qa_pairs={qa_pairs!r}"
    )


# ---------------------------------------------------------------------------
# Property 14: DOCX title paragraph present for any non-empty title
# Feature: export-formats, Property 14: DOCX title paragraph presence
# Validates: Requirements 3.2 (title in document body)
# ---------------------------------------------------------------------------


@given(
    title=_xml_safe_title,
    report_body=_xml_safe_text,
)
@settings(max_examples=50, deadline=None)
def test_property_14_docx_title_paragraph_present(title: str, report_body: str) -> None:
    """The rendered DOCX document contains the sanitized title text in at least
    one paragraph.  Control characters are stripped before writing, so the
    assertion is against the sanitized form.

    Validates: Requirements 3.2
    """
    import re as _re

    _illegal = _re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]")
    sanitized_title = _illegal.sub("", title)

    parts = DocumentParts(
        metadata=MetadataHeader(
            title=title,
            session_id="aaaaaaaa-bbbb-cccc-dddd-eeeeeeeeeeee",
            exported_at="2025-01-01T00:00:00+00:00",
            format="docx",
        ),
        report_body=report_body,
        qa_pairs=[],
    )
    result = docx_renderer.render(parts)
    doc = _DocxDocument(_BytesIO(result))

    all_text = "\n".join(p.text for p in doc.paragraphs)

    # python-docx normalises \r to \n and strips whitespace from paragraph text
    # on read-back.  Apply the same normalisation before asserting.
    normalised_title = sanitized_title.replace("\r", "\n").strip()
    if not normalised_title:
        return  # nothing meaningful to assert for whitespace-only titles
    assert normalised_title in all_text, (
        f"Normalised title {normalised_title!r} not found in document paragraphs. "
        f"Document text: {all_text[:200]!r}"
    )
