"""DOCX renderer for the export pipeline.

Pipeline: DocumentParts → python-docx Document → bytes
No LLM is involved — this is pure document assembly.
"""

from __future__ import annotations

import re
from io import BytesIO

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches, Pt

from src.export.models import DocumentParts


_HEADING_RE = re.compile(r"^(#{1,6})\s+(.+)$")
_BOLD_RE = re.compile(r"\*\*(.+?)\*\*|__(.+?)__")
_ITALIC_RE = re.compile(r"\*(.+?)\*|_(.+?)_")
_CODE_INLINE_RE = re.compile(r"`([^`]+)`")
_LINK_RE = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")
_CODE_BLOCK_RE = re.compile(r"^```(\w*)\n(.*?)^```", re.MULTILINE | re.DOTALL)
_BULLET_RE = re.compile(r"^[\-\*]\s+(.+)$")
_NUMBERED_RE = re.compile(r"^\d+\.\s+(.+)$")


def _add_formatted_run(paragraph, text: str) -> None:
    """Add text to a paragraph with inline formatting (bold, italic, code)."""
    pos = 0
    while pos < len(text):
        bold_match = _BOLD_RE.search(text, pos)
        italic_match = _ITALIC_RE.search(text, pos)
        code_match = _CODE_INLINE_RE.search(text, pos)
        link_match = _LINK_RE.search(text, pos)

        matches = [m for m in [bold_match, italic_match, code_match, link_match] if m]
        if not matches:
            run = paragraph.add_run(text[pos:])
            break

        earliest = min(matches, key=lambda m: m.start())

        if earliest.start() > pos:
            paragraph.add_run(text[pos : earliest.start()])

        if earliest == bold_match:
            content = bold_match.group(1) or bold_match.group(2)
            run = paragraph.add_run(content)
            run.bold = True
            pos = bold_match.end()
        elif earliest == italic_match:
            content = italic_match.group(1) or italic_match.group(2)
            run = paragraph.add_run(content)
            run.italic = True
            pos = italic_match.end()
        elif earliest == code_match:
            run = paragraph.add_run(code_match.group(1))
            run.font.name = "Courier New"
            run.font.size = Pt(9)
            pos = code_match.end()
        elif earliest == link_match:
            link_text = link_match.group(1)
            link_url = link_match.group(2)
            run = paragraph.add_run(f"{link_text} ({link_url})")
            run.underline = True
            pos = link_match.end()


def _parse_paragraph(text: str) -> tuple[str, str, str]:
    """Parse a paragraph and return (type, content, level).

    Types: 'heading', 'bullet', 'numbered', 'text'
    For headings, level is the heading number (1-6).
    """
    heading_match = _HEADING_RE.match(text)
    if heading_match:
        level = len(heading_match.group(1))
        return ("heading", heading_match.group(2), str(level))

    bullet_match = _BULLET_RE.match(text)
    if bullet_match:
        return ("bullet", bullet_match.group(1), "")

    numbered_match = _NUMBERED_RE.match(text)
    if numbered_match:
        return ("numbered", numbered_match.group(1), "")

    return ("text", text, "")


def _add_paragraph_with_style(doc: Document, ptype: str, content: str, level: str) -> None:
    """Add a paragraph to the document with appropriate style."""
    if ptype == "heading":
        heading_level = int(level)
        para = doc.add_heading(content, level=heading_level)
    elif ptype == "bullet":
        para = doc.add_paragraph(style="List Bullet")
        _add_formatted_run(para, content)
    elif ptype == "numbered":
        para = doc.add_paragraph(style="List Number")
        _add_formatted_run(para, content)
    else:
        para = doc.add_paragraph()
        _add_formatted_run(para, content)


def _process_markdown(doc: Document, md_text: str) -> None:
    """Process markdown text and add to document."""
    code_blocks = []
    placeholder_idx = 0

    def _code_block_replacer(match):
        nonlocal placeholder_idx
        placeholder = f"\x00CODEBLOCK{placeholder_idx}\x00"
        code_blocks.append((match.group(1) or "", match.group(2)))
        placeholder_idx += 1
        return placeholder

    md_text = _CODE_BLOCK_RE.sub(_code_block_replacer, md_text)

    paragraphs = md_text.split("\n\n")

    for para_text in paragraphs:
        if not para_text.strip():
            continue

        codeblock_match = re.match(r"\x00CODEBLOCK(\d+)\x00", para_text)
        if codeblock_match:
            idx = int(codeblock_match.group(1))
            lang, code = code_blocks[idx]
            code_para = doc.add_paragraph()
            code_para.paragraph_format.left_indent = Inches(0.5)
            if lang:
                run = code_para.add_run(f"[{lang}]\n")
                run.font.name = "Courier New"
                run.font.size = Pt(9)
                run.italic = True
            code_run = code_para.add_run(code.rstrip("\n"))
            code_run.font.name = "Courier New"
            code_run.font.size = Pt(9)
            continue

        lines = para_text.strip().split("\n")
        for line in lines:
            if not line.strip():
                continue
            ptype, content, level = _parse_paragraph(line)
            _add_paragraph_with_style(doc, ptype, content, level)


def render(parts: DocumentParts) -> bytes:
    """Render *parts* to DOCX bytes.

    Pipeline:
    1. Add metadata as document properties and title.
    2. Process report body markdown → DOCX elements.
    3. Add Q&A section if present.

    Raises:
        RenderError: if docx creation fails.
    """
    doc = Document()

    core_props = doc.core_properties
    core_props.title = parts.metadata.title
    core_props.author = "Deep Research"
    core_props.subject = f"Session {parts.metadata.session_id}"

    title = doc.add_heading(parts.metadata.title, level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    meta_para = doc.add_paragraph()
    meta_para.add_run("Session ID: ").bold = True
    meta_para.add_run(parts.metadata.session_id)
    meta_para.add_run("\nExported: ").bold = True
    meta_para.add_run(parts.metadata.exported_at)
    meta_para.paragraph_format.space_after = Pt(12)

    _process_markdown(doc, parts.report_body)

    if parts.qa_pairs:
        doc.add_heading("Q&A History", level=1)
        for pair in parts.qa_pairs:
            q_para = doc.add_paragraph()
            q_para.add_run("Q: ").bold = True
            _add_formatted_run(q_para, pair.question)

            a_para = doc.add_paragraph()
            a_para.add_run("A: ").bold = True
            _add_formatted_run(a_para, pair.answer)

            doc.add_paragraph()

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
